# -*- coding: utf-8 -*-

# imprint: a program for creating documents from data and content templates
#
# Copyright (C) 2019  Joseph R. Fox-Rabinovitz <jfoxrabinovitz at gmail dot com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# Author: Joseph Fox-Rabinovitz <jfoxrabinovitz at gmail dot com>
# Version: 13 Apr 2019: Initial Coding


"""
This module implments the base :ref:`tag-api`, as well as the
implementations of all the predefined :ref:`tag-api-builtins` and
:ref:`tag-api-refs`.
"""

from collections.abc import Mapping
import logging
from pprint import pformat
import re

from docx.enum.text import WD_BREAK
from docx.shared import Length

from haggis.files.docx import add_section, insert_toc, str2length
from haggis.latex_util import render_latex

from . import KnownError, utilities
from . import defaults
from .state import ListType


__all__ = [
    'tag_registry', 'TagDescriptor', 'ReferenceDescriptor', 'BuiltinTag',
    'get_key', 'get_handler', 'get_size', 'compute_styles', 'compute_size',
]


class _TagRegistry(dict):
    """
    Implements a limited mapping type to support the
    :py:data:`tag_registry` attribute.

    See the documentation of :py:data:`tag_registry` for a full
    description. Forbidden operations raise a :py:exc:`TypeError`.
    """
    def __delitem__(self, *args, **kwargs):
        """
        Forbidden operation.
        """
        self._delete()

    def __setitem__(self, tag, *args, **kwargs):
        """
        Register a descriptor to a specified tag name.

        Raise a :py:exc:`KeyError` if `tag` is already registered.
        """
        self._overwrite_check(tag)
        super().__setitem__(tag, *args, **kwargs)
        self._add_reference(tag, *args, **kwargs)

    def clear(self, *args, **kwargs):
        """
        Forbidden operation.
        """
        self._delete()

    def pop(self, tag, *args, **kwargs):
        """
        Operation allowed *only* if a `default` is supplied *and* the
        tag is not in the dictionary.
        """
        if tag in self:
            self._delete()
        super().pop(tag, *args, **kwargs)

    def popitem(self, *args, **kwargs):
        """
        Forbidden operation.
        """
        self._delete()

    def setdefault(self, tag, *args, **kwargs):
        """
        Operation always allowed.
        """
        self._string_check(tag)
        # Always safe: never overwrites an existing value
        item = super().setdefault(*args, **kwargs)
        self._add_reference(tag, item)
        return item

    def update(self, other):
        """
        Allowed only if none of the keys in `other` match existing
        keys.
        """
        # Checking for `keys` is how Python duck-types dicts apparently
        # This is fine because dict(other) raises the same errors as
        # update(other).
        if not hasattr(other, 'keys'):
            other = dict(other)
        for tag in other:
            self._overwrite_check(tag)
        super().update(other)
        for item in other.items():
            self._add_reference(*item)

    @property
    def referable_tags(self):
        """
        Compute a list of the names of tags that have a
        non-\ :py:obj:`None` :py:attr:`~TagDescriptor.reference`
        attribute.
        """
        return [key for key, desc in self.items()
                    if TagDescriptor.wrap(desc).reference is not None]

    def _delete(self):
        """
        Utility function for raising an error when any operation
        involving key removal is invoked.
        """
        raise TypeError('Tag deletion not supported')

    def _overwrite_check(self, tag):
        """
        Utility function for raising an error when any operation
        involving key overwrite is invoked.

        This also efficiently performs :py:meth:`_string_check` (i.e.,
        if the tag is not found).
        """
        if tag in self:
            raise KeyError(f'Tag duplication not supported: {tag!r}')
        self._string_check(tag)

    def _string_check(self, tag):
        """
        Utility function for raising an error on non-string keys.
        """
        if not isinstance(tag, str):
            raise TypeError(f'Invalid key type {type(tag).__name__}')

    def _add_reference(self, tag, item):
        """
        Invokes the item's :py:class:`ReferenceDescriptor`\ 's
        registration hook, if the item is referenceable.

        The registration hook may register additional items, so should
        be called after any requested modifications to the registry are
        complete.
        """
        item = TagDescriptor.wrap(item)
        if item.reference is not None:
            item.reference.register(self, tag, item)


#: A limited mapping type that contains all the currently registered tag
#: descriptors.
#:
#: Registering a new descriptor is as easy as doing::
#:
#:     tag_registry[name] = descriptor
#:
#: The registry is a restricted mapping type that supports adding new
#: elements only if they are not already registered. Existing elements
#: can not be deleted. Deletion operations will raise a
#: :py:exc:`TypeError`, while overwriting existing keys will raise a
#: :py:exc:`KeyError`. Aside from that, all operations supported by
#: :py:class:`dict` are allowed (including things like
#: :py:meth:`~dict.update`).
#:
#: A convenience property, :py:attr:`~_TagRegistry.referable_tags` is
#: available to compute a list of all the keys that have
#: non-\ :py:obj:`None` :py:attr:`~TagDescriptor.reference` attributes
#: in their values.
#:
#: Any tag that is :term:`referenceable` by design (has a valid
#: :py:attr:`~TagDescriptor.reference` attribute) will have the
#: :py:class:`ReferenceDescriptor`\ 's registration hook invoked after
#: the tag-proper is registered.
#:
#: The built-in tags are registered when the current module is imported.
tag_registry = _TagRegistry()


class TagDescriptor:
    """
    The basis of the tag API.

    Instances of this class contain the information required to process
    a custom tag. They *must* contain all of the attributes listed below,
    with the expected types. The elements in :py:data:`tag_registry` may
    be delegate objects that supply only part of the attibute set. In
    that case, they are wrapped in a proxy as needed at runtime, never
    up-front. The reason for this is twofold:

    1. There may be stateful objects registered for multiple tags, and
       wrapping in a proxy will not allow the tags to share state. This
       would not be a problem, except it would be unexpected behavior.
    2. Some of the attributes may be dynamic properties (or other
       descriptors). Fixing the value once would completely defeat such
       behavior.

    Creating an occasional wrapper around a delegate is not expected to
    be particularly expensive, even if it had to be done for every tag
    encountered in the XML file. On the other hand, it allows for some
    *very* flexible behaviors. At the same time, very few instances of
    wrapping should occur, since most tags will be implemented by
    extending this class and implementing it properly. The
    :py:meth:`wrap` method ensures that all extensions are passed
    through as-is.

    All the :ref:`tag-api-builtins` are instances of children of this
    class.

    .. py:attribute:: content

       A tri-state :py:class:`bool` flag indicating whether the tag is
       allowed/expected to have textual content or not. The values are
       interpreted as follows:

       None
           The tag may not have any content. It must be of the form
           ``<tag/>`` or ``<tag><otherTag>...</otherTag></tag>``.
           Anything else will raise a :term:`fatal error`. If
           :py:attr:`tags` is set to :py:obj:`False`, only the former
           form is allowed.
       False
           The tag *should* not have content, but content will not raise
           an error. A warning will be raised instead.
       True
           The tag is expected to have content, but the content may be
           empty.

       Any value is allowed in a delegate. If defined, the value will be
       converted to :py:class:`bool` if it is not :py:obj:`None`.
       Defaults to :py:obj:`None` if not defined.

    .. py:attribute:: tags

       A :py:class:`bool` indicating whether or not nested tags are
       allowed within this one.

       Any value is allowed in a delegate. If defined, the value will be
       converted to :py:class:`bool`. Defaults to :py:obj:`True` if not
       defined.

    .. py:attribute:: required

       A :py:class:`tuple` of strings containing the name of required
       tag attributes. A tag encountered without all of these attributes
       will raise an error.

       In a delegate, this may be a single string, an iterable of
       strings, :py:obj:`None` or simply omitted. Every element of an
       iterable must be a string, or a :py:exc:`TypeError` is raised
       immediately during construction. Defaults to an empty
       :py:class:`tuple` if not defined.

    .. py:attribute:: optional

       A dictionary mapping the names of optional attributes to their
       default values. Optional attributes are ones that are expected to
       be present in processing, but have sensible defaults that can be
       used, meaning that they do not have to be specified explicitly in
       the :ref:`configuration-xml`.

       In a delegate, this may be any mapping type, an iterable of
       strings, a single string, :py:obj:`None` or simply omitted. In
       the case of an iterable or individual string, all the defaults
       will be :py:obj:`None`. Iterables and mapping keys must be
       strings, or a :py:exc:`TypeError` will be raised 
       during contruction. Defaults to an empty :py:class:`dict` if not
       defined.

    .. py:attribute:: data_config

       The name of the attribute containing the data configuration name
       for the tag. This should only be provided for tags that require
       :ref:`tag-api-configuration-data`. If provided, this tag will
       automatically be added to the :py:attr:`required` sequence.

       In a delegate, this object must be an instance of :py:class:`str`
       or :py:obj:`None`. Defaults to :py:obj:`None` if not defined.

    .. py:attribute:: reference

       A :py:class:`ReferenceDescriptor` that is only present if this
       type of tag can be the target of a reference.

       Examples of referrable built-in tags are
       :ref:`xml-spec-tags-figure`, :ref:`xml-spec-tags-table` and
       sometimes :ref:`xml-spec-tags-par`. Referrable tags can have an
       optional ``role`` attribute that changes the type of reference
       they represent. See the :ref:`tag-api-references-roles`
       description for more information.

       In a delegate, this object must be an instance of
       :py:class:`ReferenceDescriptor` or :py:obj:`None`. Defaults to
       :py:obj:`None` if not defined.
    """
    def __new__(cls, *args, **kwargs):
        """
        Create an empty instance, with all required attributes set to
        default values.

        This method is provided to allow bypassing the default
        :py:meth:`__init__` in child classes. All arguments are ignored.
        """
        self = super().__new__(cls)
        self.content = None
        self.tags = True
        self.required = ()
        self.optional = {}
        self.data_config = None
        self.reference = None
        return self

    def __init__(self, delegate):
        """
        After completion, this instance has all of the required
        attributes defined in the delegate, wrapped in the required
        types.

        A reference to the delegate object is not retained. This method
        can be invoked multiple times. It updates the current descriptor
        with the attributes of the delegate, leaving undefined
        attributes in the delegate untouched.
        """
        def _string_check(iterable, label):
            """
            Verifies that all the elements of an iterable are strings.
    
            Raises :py:exc:`TypeError` with ``label`` in the message if
            the check fails.
            """
            for item in iterable:
                if not isinstance(item, str):
                    raise TypeError(f'{label} tag names must be strings: '
                                    f'found {type(item).__name__}')

        content = getattr(delegate, 'content', self.content)
        self.content = content if content is None else bool(content)

        self.tags = getattr(delegate, 'tags', self.tags)

        required = getattr(delegate, 'required', self.required)
        if required is None:
            required = ()
        elif isinstance(required, str):
            required = (required,)
        else:
            # Implicit iterable check
            required = tuple(required)
        _string_check(required, 'required')
        self.required = required

        optional = getattr(delegate, 'optional', self.optional)
        if optional is None:
            optional = {}
        elif isinstance(optional, str):
            optional = {optional: None}
        elif not isinstance(optional, dict):
            if isinstance(optional, Mapping):
                optional = dict(optional)
            else:
                # Implicit iterable check
                optional = dict.fromkeys(optional)
        _string_check(optional, 'optional')
        self.optional = optional

        start = getattr(delegate, 'start', None)
        if start is not None:
            # Override the non-data method descriptor with a bound method
            self.start = start

        end = getattr(delegate, 'end', None)
        if end is not None:
            # Override the non-data method descriptor with a bound method
            self.end = end

        data_config = getattr(delegate, 'data_config', self.data_config)
        if data_config is not None:
            data_config = str(data_config)
            if data_config not in required:
                self.required += (data_config,)
        self.data_config = data_config

        reference = getattr(delegate, 'reference', self.reference)
        if reference is not None:
            if not isinstance(reference, ReferenceDescriptor):
                raise TypeError('Reference may not be a '
                                f'{type(reference).__name__}')
            for attr in reference.identifiers:
                if attr not in self.required and attr not in self.optional:
                    raise ValueError(
                        'All reference identifiers must be registered '
                        f'attributes. Found {attr!r}.'
                    )
        self.reference = reference

    def start(self, state, name, attr, *args):
        """
        Each descriptor should provide a method with this signature to
        process opening tags.

        If implemented, this method **must** accept the
        :ref:`tag-api-engine-state`, a tag name and a :py:class:`dict`
        of attributes. Normally, the tag name is ignored since a separate
        descriptor is registered for each tag.

        Descriptors that have a non-\ :py:obj:`None`
        :py:attr:`data_config` attribute set will receive an additional
        argument containing the :ref:`tag-api-configuration-data`.

        The default implementation just logs itself.
        """
        state.log(logging.TRACE, 'Opening tag <%s>: No processing required',
                  name)

    def end(self, state, name, attr, *args):
        """
        Each descriptor should provide a method with this signature to
        process closing tags.

        If implemented, this method **must** accept the
        :ref:`tag-api-engine-state`, a tag name and a :py:class:`dict`
        of attributes. Normally, the tag name is ignored since a separate
        descriptor is registered for each tag. The attributes are the
        same as those passed to :py:meth:`start`, barring any
        modifications made *in* :py:meth:`start`.

        Descriptors that have a non-\ :py:obj:`None`
        :py:attr:`data_config` attribute set will receive an additional
        argument containing the :ref:`tag-api-configuration-data`.

        The default implementation just logs itself.
        """
        state.log(logging.TRACE, 'Closing tag <%s>: No processing required',
                  name)

    @classmethod
    def wrap(cls, desc):
        """
        Construct a proxy from the descriptor if it isn't already one.

        This method is provided so that when :py:class:`TagDescriptor`
        objects are implemented properly up front, they do not need to
        be wrapped in an additional layer.

        If the input is a delegate, the return value will always be of
        the type that this method was invoked on. However, the type
        check will always be done agains the base
        :py:class:`TagDescriptor` class.
        """
        return desc if isinstance(desc, __class__) else cls(desc)


class ReferenceDescriptor:
    """
    Defines the process for creating :ref:`tag-api-references` and using
    them through the appropriate tag.

    References are made by processing the :ref:`configuration-xml` and
    mapping out any :term:`referenceable` tags using the
    :py:meth:`start` and :py:meth:`end` methods. In the default
    implementation, the reference text is created by the
    :py:meth:`make_reference` method, invoked from :py:meth:`end`.

    :py:meth:`start` and :py:meth:`end` return a boolean value to allow
    custom tags to be processed selectively. A return value of
    :py:obj:`False` from *either* method means that that the specific
    instance of the tag being processed is not a valid reference target.
    Normally both methods always return :py:obj:`True`, but for the
    builtin :ref:`xml-spec-tags-par` tag, for example, an exception must
    be made.

    References are placed into the document by a special
    :py:class:`TagDescriptor`, which is generally registered along with
    the parent tag that contains a :py:class:`ReferenceDescriptor` using
    the :py:meth:`register` method.

    Current references are purely textual, rather having a dynamic field
    assigned to them. This is still a
    :ref:`work in progress <future-work-anchors>`.

    .. py:attribute:: prefix

       The prefix that normally gets prepended to the reference text.
       Used by :py:meth:`make_reference` to construct the output string.
       Extensions are welcome to ignore this attribute.

    .. py:attribute:: identifiers

       A string or iterable of strings that lists the attributes that
       are used to identify target for this reference type. The
       attribute may be either required or optional for the target
       tag, but it must be recognized either way. This attribute is used
       to check for attributes on tags with a non-default
       :ref:`xml-spec-attributes-role`. Defaults to ``'id'``.
    """
    def __init__(self, prefix, identifiers='id'):
        if isinstance(identifiers, str):
            identifiers = (identifiers,)
        else:
            identifiers = tuple(identifiers)
            for attr in identifiers:
                if not isinstance(attr, str):
                    raise ValueError(
                        'All reference identifier attributes must be strings: '
                        f'found {type(attr).__name__}'
                    )
        self.__dict__.update(prefix=str(prefix), identifiers=identifiers)

    @property
    def prefix(self):
        """
        Ensure that :py:attr:`prefix` is read-only.
        """
        return self.__dict__['prefix']

    @property
    def identifiers(self):
        """
        Ensure that :py:attr:`identifiers` is read-only.
        """
        return self.__dict__['identifiers']

    def start(self, state, name, role, attr):
        """
        Process the opening tag for a referencable tag.

        The default is to log the tag and its role.

        Returns :py:obj:`True` if the tag is a potential reference
        target, :py:obj:`False` if it is definitely not.
        """
        state.log(logging.TRACE, 'Opening <%s> with role %r', name, role)
        return True

    def end(self, state, name, role, attr):
        """
        Process the closing tag for a referencable tag.

        The default is to add the reference to the appropriate map in
        :py:attr:`~imprint.core.state.ReferenceState.references` by ID,
        based on the ``role``, and log the operation. The attribute
        ``id`` is required.

        The actual reference is created by :py:meth:`make_reference`.

        Returns :py:obj:`True` if the tag is definitely a reference
        target, :py:obj:`False` if not.
        """
        computed_role = name if role is None else role
        reference = self.make_reference(state, computed_role, attr)
        self.set_reference(state, computed_role, 'id', attr['id'], reference)
        state.log(logging.DEBUG, 'Closed <%s> with role %r: '
                  'added reference "%s"', name, role, reference)
        return True

    def register(self, registry, name, descriptor):
        """
        A registration hook that is invoked when the parent
        :py:class:`TagDescriptor` is registered.

        The default implementation registers an additional
        :py:class:`TagDescriptor` under the name ``name + '-ref'``,
        which replaces the ``<name-ref/>`` tag with the formatted
        reference. See :py:class:`ReferenceProcessor`.

        Parameters
        ----------
        registry :
            The tag registry that the parent :py:class:`TagDescriptor`
            is being inserted into. See :py:data:`tag_registry` for
            details on the interface.
        name : str
            The name under which the *parent* tag is being registered.
        descriptor :
            The parent object being registered, not necessarily a
            :py:class:`TagDescriptor`. The :py:meth:`TagDescriptor.wrap`
            method can be used to retreive the corresponding
            :py:class:`TagDescriptor` if necessary.
        """
        registry[name + '-ref'] = ReferenceProcessor.get_instance()

    def make_reference(self, state, role, attr):
        """
        Returns a string refering to the specified tag in the specified
        role.

        Keep in mind that the :py:class:`ReferenceDescriptor` is
        selected based on the role, not necessarily the tag name.
        Therefore, the ``role`` argument should always be the "computed"
        role: the name of the tag should be overriden by the value of
        the attribute, if it was specified.
        """
        state.item_counters[role] += 1
        state.log(logging.TRACE, 'Creating new reference for %r, #%d',
                  role, state.item_counters[role])
        return state.format_heading(
            prefix=self.prefix, sep=defaults.reference_level_sep,
            suffix_sep=defaults.reference_sep, suffix=state.item_counters[role]
        )

    def set_reference(self, state, role, attribute, key, reference,
                      duplicates=False):
        """
        Check that the reference identified by ``key`` does not already
        exist and set it.

        Duplicate reference targets cause an error, unless
        ``duplicates`` is :py:obj:`True`, in which case a warning is
        logged and the new value is discarded.
        """
        combined = (role, attribute, key)
        if combined in state.references:
            if duplicates:
                state.log(
                    logging.WARN, 'Duplicate %s %s %r found for "%s" and now '
                    '"%s"', *combined, state.references[combined], reference
                )
            else:
                msg = 'Duplicate %s %s %r'
                state.log(logging.ERROR, msg, *combined)
                raise ValueError(msg % combined)
        else:
            state.log(logging.INFO, 'Adding new reference '
                      '[role=%s, attr=%s, key=%r] = %r', *combined, reference)
            state.references[combined] = reference


class SegmentReferenceDescriptor(ReferenceDescriptor):
    """
    Extension of :py:class:`ReferenceDescriptor` to accumulate heading
    text and allow references through the ``title`` attribute.

    Used by :ref:`xml-spec-tags-par` tags to create heading references.

    .. py:attribute:: heading_style_name_pattern

       A class-level :py:ref:`regular expression <re-objects>` for
       identifying the :ref:`xml-spec-tags-par` tags that represent
       referenceable headings.
    """
    heading_style_name_pattern = re.compile(r'heading (\d+)')

    def start(self, state, name, role, attr):
        """
        Start accumulating content in addition to the default logging.

        If an actual :ref:`xml-spec-tags-par` tag is encountered (as
        opposed to a tag playing that :ref:`xml-spec-attributes-role`),
        and the heading matches ``Heading \\d+``, the current heading
        is incremented in the state.

        If any heading tag, or any tag with ``role="par"`` is
        encountered, a new reference will be created. Non-heading
        paragraphs with no explicit role are non-referenceable. A
        non-heading paragraph can be made referenceable by explicitly
        setting the :ref:`xml-spec-attributes-role`.

        Keep in mind that the title for a segment reference is
        accumulated from **all** the text in the paragraph. Use
        carefully with non-default tags.
        """
        style = attr['style'] or defaults.paragraph_style
        if name == 'par':
            match = self.heading_style_name_pattern.fullmatch(style.lower())
            if match:
                try:
                    level = int(match.group(1))
                except (ValueError, TypeError) as e:
                    state.log(logging.INFO, 'Spurious invalid heading level %r',
                              level)
                    return False
                state.increment_heading(level)
            elif role is None:
                return False

        state.start_content()
        return super().start(state, name, role, attr)

    def end(self, state, name, role, attr):
        """
        Create a dual reference based on the title and optional ID
        in addition to the default logging.
        """
        computed_role = name if role is None else role
        title = state.end_content()
        reference = self.make_reference(state, computed_role, attr, title)
        if attr['id']:
            self.set_reference(state, computed_role, 'id', attr['id'],
                               reference)
            where = 'to ID and TITLE maps'
        else:
            where = 'to TITLE map only'
        self.set_reference(state, computed_role, 'title', title,
                           reference, duplicates=True)
        state.log(logging.DEBUG, 'Closed <%s> with role %r: added '
                  'reference "%s" %s', name, role, reference, where)
        return True

    def make_reference(self, state, role, attr, title):
        """
        Add the section heading to the usual reference text.
        """
        ref = state.format_heading(
            prefix=self.prefix, sep=defaults.reference_level_sep,
            suffix_sep=defaults.reference_sep, suffix=None
        )
        if title:
            ref = f'{ref}: {title}'
        return ref

    def register(self, registry, name, descriptor):
        """
        Register a :py:class:`SegmentRefProcessor` for the
        :ref:`xml-spec-tags-segment-ref` tag.

        This registration hook uses a fixed name, so can only be called
        once.
        """
        registry['segment-ref'] = SegmentRefProcessor.get_instance()


class BuiltinTag(TagDescriptor):
    """
    The base class of all the built-in :py:class:`TagDescriptor`
    implementations.

    Custom tag implementations are welcome to use this class as a base
    instead of a raw :py:class:`TagDescriptor`.
    """
    def __init__(self, delegate=None, **kwargs):
        """
        Updates the required fields with the keywords that are passed
        in.

        If no delegate object (or :py:obj:`None`) is supplied, bypass
        the default constructor (see :py:meth:`TagDescriptor.__new__`).
        `kwargs` will override any defaults and attributes set by a
        delegate.
        """
        if delegate is not None:
            super().__init__(delegate)
        self.__dict__.update(kwargs)


def get_key(key, attr, data, sentinel=None, default=None):
    """
    Resolve the value of ``key`` with respect to ``attr``, but with
    the option to override by the `data` configuration dictionary.

    If the final value is `sentinel`, return `default` instead.
    Return `default` if `key` is missing entirely as well. Both
    `attr` and `data` must be mapping types that support
    a `get` method.
    """
    value = data.get(key, attr.get(key, sentinel))
    return default if value == sentinel else value


def get_size(attr, data, key='size'):
    """
    Convert a string, number or pre-constructed size to a
    :py:class:`docx.shared.Length` object, using :py:func:`get_key`
    for value resolution.

    Common options for ``key`` are ``'width'`` and ``'height'``.

    Valid units suffixes are ``"``, ``in``, ``cm``, ``mm``, ``pt``,
    ``emu``, ``twip``. Default when no units are specified is inches
    (``"``).
    """
    size = get_key(key, attr, data)
    if size is not None and not isinstance(size, Length):
        size = str2length(size)
    return size


def get_handler(tag, id, attr, data, logger, key='handler'):
    """
    Retrieve and load the :ref:`handler <plugins-handlers>` for the
    specified attribute mapping and data configuration.

    If the handler can not be found, a detailed exception is logged and
    a :py:exc:`~imprint.core.KnownError` is raised.
    """
    handler_name = get_key(key, attr, data)
    try:
        handler = utilities.get_handler(handler_name)
    except TypeError as e:
        raise KnownError(f'Unable to find handler "{handler_name}" '
                         f'for {tag} "{id}"') from e
    logger(logging.TRACE, 'Found %r for %r "%s"', handler, tag, id)
    return handler


def get_and_run_handler(tag, id, attr, data, logger, args, kwargs=None,
                        key='handler'):
    """
    Load and run the :ref:`handler <plugins-handlers>` for the
    specified attribute mapping and data configuration.

    If the handler can not be found, a detailed exception is logged, as
    with :py:func:`get_handler`.

    *All* exceptions that occur during execution are converted into
    :py:exc:`~imprint.core.KnownError`.
    """
    handler = get_handler(tag, id, attr, data, logger, key)
    if kwargs is None:
        kwargs = {}
    try:
        return handler(*args, **kwargs)
    except Exception as e:
        raise KnownError(f'Error processing handler for {tag} "{id}": '
                         f'{str(e)}') from e


def compute_styles(attr, data, defaults):
    """
    Compute the required styles based on `attr` and `data`
    configurations.

    Style keys are taken from the keys of `defaults`, while values
    provide the fallback names used if the keys do not appear in either
    `attr` or `data`. Similarly named keys in `data` will override ones
    in ``attr``.
    """
    return {
        style: get_key(style, attr, data, default=default)
            for style, default in defaults.items()
    }


def compute_size(tag, attr, data, logger, default_width=None,
                 width_key='width', height_key='height'):
    """
    Create a dictionary with keys ``width`` and ``height`` and values
    that are instances of :py:class:`docx.shared.Length`.

    Values are resolved according to the rules of :py:func:`get_key`,
    with ``width_key`` and ``height_key`` as the inputs. String values
    may contain units, and will be parsed according to
    :py:func:`get_size`.

    If neither key is present in either configuration (or present but
    set to :py:obj:`None`), set the the width to `default_width`. If
    that is :py:obj:`None` as well, return an empty dictionary.
    """
    width = get_size(attr, data, width_key)
    height = get_size(attr, data, height_key)
    if width is None and height is None:
        if default_width is None:
            message = '%s size not set explicitly.'
            ret = {}
        else:
            message = '%s size not set explicitly. Default width will be used.'
            ret = {'width': default_width}
        args = tag.title(),
    elif width is None:
        message, args = 'Setting custom %s height to %s"', (tag, height.inches)
        ret = {'height': height}
    elif height is None:
        message, args = 'Setting custom %s width to %s"', (tag, width.inches)
        ret = {'width': width}
    else:
        message = 'Fixed-size figure height to W%s" x H%s"'
        args = (width.inches, height.inches)
        ret = {'width': width, 'height': height}

    logger(logging.TRACE, message, *args)
    return ret


class BreakTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-break` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('tags', False)
        super().__init__(delegate, **kwargs)

    def end(self, state, name, attr):
        """
        Insert a page break into the document.
        """
        if 'run' in state:
            state.flush_run()
            state.run.add_break(WD_BREAK.PAGE)
        else:
            state.doc.add_page_break()
        state.log(logging.TRACE, 'End of page break')


class ExprTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-expr` tag.

    .. warning::

       This descriptor uses :py:func:`eval` to execute arbitrary code
       and assign it to a new keyword. Use with extreme caution!
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', True)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('required', ('name',))
        kwargs.setdefault('optional', {'imports': ''})
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Begin a new expression.

        This just pushes a new
        :py:attr:`~imprint.core.state.EngineState.content_stack` entry
        in the `state`. All content until the closing tag will be
        evaluated as a set of Python statements.
        """
        state.log(logging.TRACE, 'Beginning expression evaluation')
        state.push_content_stack(flush=False)

    def end(self, state, name, attr):
        """
        Evaluate the expression found inside the tag, and add a new
        entry to the ``state``\ 's
        :py:attr:`~imprint.core.state.EngineState.keywords`.

        The :py:attr:`~imprint.core.state.EngineState.content_stack` will
        be popped.

        All errors in importing and evaluation will be propagated up and
        will terminate the parser.
        """
        keyword = attr['name']
        raw = state.content.getvalue()
        imports = {}
        for item in attr['imports'].split():
            name = item.split('.')[0]
            try:
                imports[name] = __import__(item)
            except ImportError:
                state.log(logging.ERROR, 'Error importing "%s" for expression '
                          '"%s": The following evaluation is likely to crash.',
                          item, keyword, exc_info=True)
                raise
        state.log(logging.DEBUG, 'Evaluating expression for "%s" with imports '
                  '"%s": %s', keyword, attr['imports'], raw)
        try:
            value = eval(raw, imports, state.keywords)
        except:
            state.log(logging.ERROR, 'Error evaluating expression for "%s"',
                      keyword)
            raise
        if keyword in state.keywords:
            state.log(logging.WARNING, 'Keyword "%s" already defined. '
                      'Overwriting value "%s" with "%s"', keyword,
                       state.keywords[keyword], value)
        state.log(logging.DEBUG, 'Setting "%s"="%s"', keyword, value)
        state.keywords[keyword] = value
        state.pop_content_stack()
        state.log(logging.TRACE, 'End of expression evaluation')


class FigureTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-figure` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', False)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('required', ('id', 'handler'))
        kwargs.setdefault('optional', {
            'style': None,
            'pstyle': None,
            'width': None,
            'height': None,
        })
        kwargs.setdefault('reference', ReferenceDescriptor('Figure'))
        kwargs.setdefault('data_config', 'id')
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr, config):
        """
        Just log the tag.
        """
        state.log(
            logging.DEBUG, 'Starting %s with ID="%s", Handler="%s", '
            'Config=%s', name, attr['id'], attr['handler'],
            pformat(config, indent=4)
        )
        state.push_content_stack(flush=True)

    def end(self, state, name, attr, config):
        """
        Generate and insert a figure based on the selected handler.

        Figures can appear in a run, a paragraph, or on their own.
        """
        id = attr[self.data_config]

        styles = compute_styles(attr, config, {
            'style': defaults.figure_run_style,
            'pstyle': defaults.figure_paragraph_style
        })
        size = compute_size(name, attr, config, logger=state.log,
                            default_width=defaults.figure_width)

        if state.keywords.get('log_images', False):
            # If dumping to file, generate text file name of figure
            figure = state.image_log_name(id)
            figure = get_and_run_handler(
                name, id, attr, config, logger=state.log,
                args=(config, state.keywords, figure)
            )
            if figure is not None:
                state.log(logging.TRACE, 'Saved figure to %s', figure)
        else:
            # If dumping to memory, generate figure in memory
            figure = get_and_run_handler(
                name, id, attr, config, logger=state.log,
                args=(config, state.keywords)
            )
            state.log(logging.TRACE, 'Generated figure in-memory')

        if figure is None:
            state.log(logging.WARNING, 'Unable to create image for '
                      '%s "%s" (no error)', name, id)
            raise KnownError(f'No error in {name} generator {id}')

        state.log(logging.TRACE, 'Successfully generated %s %r', name, id)
        state.insert_picture(figure, **size, **styles)
        state.log(logging.INFO, 'Successfully inserted image for %s %r',
                  name, id)

        state.pop_content_stack()
        state.log(logging.TRACE, 'Figure generation ends')


class KwdTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-kwd` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', False)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('required', ('name',))
        kwargs.setdefault('optional', {'format': ''})
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Find the value of the keyword in the state's
        :py:attr:`~imprint.core.state.EngineState.keywords` and place it
        into the current
        :py:attr:`~imprint.core.state.EngineState.content`.

        If the keyword is not found, a :py:exc:`KeyError` will be
        raised. If the tag has a ``format`` attribute, it is interpreted
        as a :token:`format_spec`, and used to convert the value. If the
        attribute is not present, the value is converted with a simple
        call to :py:class:`str`.
        """
        state.log(logging.TRACE, 'Beginning keyword replacement')
        keyword = attr['name']
        value = state.keywords[keyword]  # Propagate a KeyError here
        state.log(logging.DEBUG, 'Performing keyword replacement for '
                  '"%s"="%s"', keyword, value)
        string = f'{value:{attr["format"]}}'
        state.content.write(string)


class LatexTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-latex` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', True)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('optional', {
            'style': None,
            'pstyle': None,
            'size': None,
            'dpi': defaults.latex_dpi,
            'format': defaults.latex_format,
        })
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Begin a new LaTeX formula.

        Just push a new
        :py:attr:`~imprint.core.state.EngineState.content_stack` entry
        into `state`. All content until the closing tag is evaluated as
        a LaTeX document.
        """
        state.log(logging.TRACE, 'Beginning latex formula')
        state.push_content_stack(flush=True)
        # Get a fresh stack even if the run already made that happen.
        # Avoid conflict with spurious text.

    def end(self, state, name, attr):
        """
        Convert the equation in the text of the current tag into an
        image using :py:meth:`haggis.latex_util.render_latex`, and
        insert the image into the parent tag.

        The parent can be a run or a paragraph. If the requested run
        style does not match the current run, the current run will be
        interrupted by a run containing a new picture with the requested
        style, and resumed afterwards. If there is no run to begin with,
        a new run will be created, but not stored in the
        :py:attr:`~imprint.core.state.EngineState.run` attribute of the
        state.

        Formulas are rendered at 96dpi in JPEG format by default.
        """
        text = state.get_content()
        state.log(logging.DEBUG, 'Appending latex formula "%s"', text)

        args = {'format': attr['format'], 'dpi': attr['dpi']}
        size = attr['size']
        if size is not None:
            args['fontsize'] = int(size)

        if state.keywords.get('log_images', False):
            # If dumping to file, generate text file name of figure
            if 'latex_count' in state:
                state.latex_count += 1
            else:
                state.latex_count = 1
            pic = state.image_log_name(f'LaTeX{state.latex_count:04d}',
                                       f'.{attr["format"]}')
            render_latex(text, pic, **args)
            state.log(logging.TRACE, 'Saved latex to %s', pic)
        else:
            # If dumping to memory, generate figure in memory
            pic = render_latex(text, **args)

        styles = compute_styles(attr, {}, {
            'style': defaults.latex_run_style,
            'pstyle': defaults.latex_paragraph_style,
        })

        if 'run' in state:
            state.run.add_text(' ')

        # pop the content stack before inserting the picture to clear the
        # buffer for the run.
        state.pop_content_stack()

        state.log(logging.TRACE, 'Successfully generated %s image', name)
        state.insert_picture(pic, **styles)
        state.log(logging.INFO, 'Successfully inserted %s image', name)

        state.log(logging.TRACE, 'End of latex formula')


class NTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-n` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('tags', False)
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Add a line break to the current run.

        If not inside a run, append the break to the last run. Make a
        new run only at the start of a paragraph. Ignore with a warning
        outside of a paragraph.
        """
        state.log(logging.TRACE, 'Found newline')
        if 'paragraph' in state:
            if 'run' in state:
                state.log(logging.TRACE, 'Adding to current run')
                state.flush_run()
                state.content.leading_space = False
                run = state.run
            elif state.paragraph.runs:
                state.log(logging.TRACE, 'Adding to previous run')
                run = state.paragraph.runs[-1]
            else:
                state.log(logging.TRACE,
                          'Adding new run to start of paragraph')
                # Do not retain a reference to this run
                run = state.paragraph.add_run(style=defaults.run_style)
            run.add_break(WD_BREAK.LINE)
        else:
            state.log(logging.WARNING, 'Spurious <%s> tag outside of <par>',
                      name)


class ParTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-par` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', False)
        kwargs.setdefault('optional', {
            'style': None,
            'id': None,
            'list': None,
            'list-level': None,
        })
        kwargs.setdefault('reference', SegmentReferenceDescriptor('Section'))
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Terminate any existing paragraph, flush all text and start a
        new paragraph.

        If the new paragraph is a list item, add the necessary metadata
        to it.

        Issue a warning if an existing paragraph is found.
        """
        if 'paragraph' in state:
            state.log(logging.WARNING, 'Nested paragraphs are forbidden. '
                      'Previous paragraph will be terminated.')
            state.end_paragraph(name)

        list_type = self.check_list(state, attr)
        style = self.compute_paragraph_style(state, attr, list_type)
        state.log(logging.TRACE, 'Beginning paragraph with style "%s"', style)
        state.paragraph = state.doc.add_paragraph(style=style)
        if list_type is not None:
            state.log(logging.TRACE, 'Paragraph is a list item of type %r',
                      list_type.value)
        state.number_paragraph(list_type, attr['list-level'])
        

    def end(self, state, name, attr):
        """
        Terminate the current paragraph.

        See :py:meth:`~imprint.core.state.EngineState.end_paragraph` in
        :py:class:`~imprint.core.state.EngineState`.
        """
        state.end_paragraph()

    def check_list(self, state, attr):
        """
        Validate the ``list`` attribute that is found.

        Log an error if the attribute is invalid, but do not terminate
        processing. The attribute is simply ignored if the list is
        neither numbered, bulleted nor continued.

        Return the type normalized to a
        :py:class:`~imprint.core.state.ListType`, or :py:obj:`None` if
        not a list item. If the type is valid, and ``list-level`` is
        set, it is converted to an integer.
        """
        list_type = attr['list']

        if list_type is None:
            state.log(logging.TRACE, 'Not a list paragraph')
            return None

        try:
            list_type = ListType.convert(list_type)
        except ValueError as e:
            state.log(logging.ERROR, 'Unable to parse list type %r',
                      list_type, exc_info=True)
            return None

        level = attr['list-level']
        if level is None:
            label = '<default>'
        else:
            attr['list-level'] = int(level)
            label = level

        state.log(logging.TRACE, '%s list paragraph, level = %s',
                  list_type.value.title(), label)

        return list_type

    def compute_paragraph_style(self, state, attr, list_type):
        """
        Compute the paragraph style based on whether an explicit style
        is set in the attributes, and whether or not the paragraph is a
        list.

            1. If an explicit style is requested, return it.
               Otherwise:
            2. If the paragraph is *not* a list, return the default
               paragraph style. Otherwise:
            3. If the previous paragraph is a list item in the same list
               (i.e., the current ``list-level`` attribute is non-zero),
               return the style of the previous paragraph. Otherwise:
            4. Return the default list item style.

        Parameters
        ----------
        state : ~imprint.core.state.EngineState
            The state is used to check for the previous item's style in
            case #3.
        attr : dict
            The tag attributes, used to check for an explicitly set
            ``style`` as well as for a style reset with
            ``list-level = 0``.
        list_type : ~imprint.core.state.ListType or None
            The type of the list, if a list at all, as returned by
            :py:meth:`check_list`.
        """
        style = attr['style']
        if style is not None:
            return style

        if list_type is None:
            return defaults.paragraph_style

        if list_type is ListType.CONTINUED:
            return state.last_list_item.style.name

        if list_type is ListType.NUMBERED:
            return defaults.numbered_list_paragraph_style

        if list_type is ListType.BULLETED:
            return defaults.bulleted_list_paragraph_style

        # Catchall that "should never happen"
        raise ValueError(f'Unsupported ListType {list_type!r}')


class RunTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-run` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', True)
        kwargs.setdefault('optional', {'style': defaults.run_style,})
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Create a new run, ensuring that there is a paragraph to go with
        it.

        Creating a run outside a paragraph raises a warning and creates
        a paragraph with a default style. See
        :py:meth:`imprint.core.state.EngineState.new_run`.
        """
        style = attr['style']
        state.log(logging.TRACE, 'Beginning run with style "%s"', style)
        state.new_run(name, style)

    def end(self, state, name, attr):
        """
        Place any remaining text into the current run, and remove
        :py:attr:`~imprint.core.state.EngineState.run` attribute of
        ``state``.
        """
        state.flush_run()
        del state.run
        state.log(logging.TRACE, 'Ended run')


class SectionTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-section` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('tags', False)
        kwargs.setdefault('optional', {'orientation': 'portrait'})
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Begin a new section in the document, optionally altering the
        page orientation.
        """
        state.log(logging.TRACE, 'Start of new document section')
        orientation = attr['orientation']
        with state.interrupt_paragraph(warn=name):
            add_section(self.doc, orientation)
        state.log(logging.INFO, 'Section started with "%s" orientation',
                  orientation)


class SkipTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-skip` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', True)
        super().__init__(delegate, **kwargs)


class StringTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-string` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', False)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('required', ('id', 'handler'))
        kwargs.setdefault('data_config', 'id')
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr, config):
        """
        Just log the tag.
        """
        state.log(logging.DEBUG, 'Starting String with ID="%s", Handler="%s"',
                  attr['id'], attr['handler'])
        state.push_content_stack(flush=True)

    def end(self, state, name, attr, config):
        """
        Generate a string based on the appropriate handler.

        If the ``log_images`` key is set to a truthy value in
        ``state``\ .
        :py:attr:`~imprint.core.state.EngineState.keywords`,
        the content will also be dumped to a file.
        """
        # Ignore contents of tag completely. Reinstate previous content so that
        # a string will be added to it.
        state.pop_content_stack()

        id = attr[self.data_config]
        string = get_and_run_handler(name, id, attr, config, logger=state.log,
                                     args=(config, state.keywords))

        if string is None:
            state.log(logging.WARNING, 'Unable to create text for '
                      '%s "%s" (no error)', name, id)
            raise KnownError(f'No error in {name} generator {id}')

        state.log(logging.DEBUG, 'Successfully generated %s %r: "%s"',
                  name, id, string)

        if state.keywords.get('log_images', False):
            fname = state.image_log_name(id, '.txt')
            with open(fname, 'wt') as file:
                file.write(string)
            state.log(logging.DEBUG, 'Saved %s %r to "%s"', name, id, fname)

        state.content.write(string)
        state.log(logging.INFO, 'Successfully inserted %s %r', name, id)

        state.log(logging.TRACE, 'String generation ends')


class TableTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-table` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', False)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('required', ('id', 'handler'))
        kwargs.setdefault('optional', {'style': None})
        kwargs.setdefault('reference', ReferenceDescriptor('Table'))
        kwargs.setdefault('data_config', 'id')
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr, config):
        """
        Just log the tag.
        """
        state.log(logging.DEBUG, 'Starting Table with ID="%s", Handler="%s"',
                  attr['id'], attr['handler'])
        state.push_content_stack(flush=True)

    def end(self, state, name, attr, config):
        """
        Generate and inserts a table based on the selected handler.

        The handler creates the table directly in the document (unlike
        for figures, where only the final product is inserted). Any
        error that occurs mid-processing leaves a stub table in the
        document in addition to the automatically-inserted alt-text.

        Tables appear on their own, outside any paragraph or run, so if
        a table is nested in a run or paragraph, a warning will be
        issued. Any interrupted run or paragraph resumes after the table
        with their prior styles.
        """
        id = attr[self.data_config]
        style = get_key('style', attr, config, default=defaults.table_style)
        image_log = state.image_log_name(id) if \
            state.keywords.get('log_images', False) else None
        with state.interrupt_paragraph(warn=name):
            get_and_run_handler(
                name, id, attr, config, logger=state.log,
                args=(config, state.keywords, state.doc, style),
                kwargs={'image_log_name': image_log},
            )
        state.pop_content_stack()
        state.log(logging.TRACE, 'Table generation ends')


class TocTag(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-toc` tag.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', True)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('optional', {
            'min': defaults.toc_min_level,
            'max': defaults.toc_max_level,
            'style': defaults.toc_title_style,
        })
        super().__init__(delegate, **kwargs)

    def start(self, state, name, attr):
        """
        Create a new TOC.

        Log a warning if the tag appears within a paragraph. Truncate
        the paragraph, and resum with the prior style. The same happens
        to the current run, if there is one.
        """
        state.log(logging.TRACE, 'Stubbing out TOC')
        state.log(logging.DEBUG, 'TOC will cover headings from %d to %d',
                  attr['min'], attr['max'])

        # TOC always has a private content stack
        state.push_content_stack(flush=True)

    def end(self, state, name, attr):
        """
        Terminate and insert the TOC.

        Gather any text that has been acquired into the heading, which
        will be a separate pargraph preceding the TOC.

        If the TOC interrupted an existing paragraph, a new paragraph
        will be resumed with the same style as the original. If a run
        style is present as well, a run will be recreated too.
        """
        with state.interrupt_paragraph(name):
            title = state.get_content()
            if title:
                style = attr['style']
                state.log(logging.DEBUG, 'Title of TOC=%s (using style=%r)',
                          title, style)
                state.doc.add_paragraph(style=style).text = title
            insert_toc(state.doc, attr['min'], attr['max'])
            state.pop_content_stack()
        state.log(logging.TRACE, 'End of TOC stub')


class ReferenceProcessor(BuiltinTag):
    """
    Implements the :ref:`xml-spec-tags-figure-ref` and
    :ref:`xml-spec-tags-table-ref` tags.

    This processor is not registered explicitly. It gets added by all of
    the target tags that use it as part of their registration process.
    Registering this processor under a name that does not end in
    ``'-ref'`` will lead to a runtime error in :py:meth:`resolve`.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('content', None)
        kwargs.setdefault('tags', False)
        kwargs.setdefault('required', ('id',))
        super().__init__(delegate, **kwargs)

    def end(self, state, name, attr):
        """
        Insert a string with the specified reference into the current
        :py:attr:`~imprint.core.state.EngineState.content`.
        """
        state.log(logging.TRACE, 'Start of %s reference resolution', name)
        element = self.resolve(state, name, attr)
        state.content.write(str(element))

    def resolve(self, state, name, attr):
        """
        Overridable operation for fetching and logging the reference
        that is to be inserted.

        The default is to look up the reference by ``'id'`` in the
        :py:class:`imprint.core.state.EngineState`\ 's.
        :py:attr:`~imprint.core.state.EngineState.references`.

        Used by the default implementation of :py:meth:`end`.
        """
        cname = name.lower()
        if not cname.endswith('-ref'):
            raise ValueError(f'Unable to resolve reference {name}: '
                             'does not end in "-ref"')
        tag = cname[:-4]
        id = attr['id']
        element = state.references[tag, 'id', id]
        state.log(logging.INFO, 'Resolved <%s id="%s"> to "%s"',
                  name, id, element)
        return element

    @classmethod
    def get_instance(cls):
        """
        Returns a quasi-singleton instance of the current class.

        This instance is not exposed directly, but it is registered by
        the built-in referencable tags.
        """
        if '_singleton' not in cls.__dict__:
            cls._singleton = cls()
        return cls._singleton


class SegmentRefProcessor(ReferenceProcessor):
    """
    Implements the :ref:`xml-spec-tags-segment-ref` tag.

    This is a special case of :py:class:`ReferenceProcessor` that allows
    access by both ``title`` and ``id``. It's references always resolve
    to a :ref:`xml-spec-tags-par` tag, or a tag playing that
    :ref:`xml-spec-attributes-role`.
    """
    def __init__(self, delegate=None, **kwargs):
        kwargs.setdefault('required', ())
        kwargs.setdefault('optional', {'id': None, 'title': None})
        super().__init__(delegate, **kwargs)

    def resolve(self, state, name, attr):
        """
        Resolve a segment reference be either text or ID.

        Either the ``id`` or ``title`` tag attribute must be present. If
        both are present, they must resolve to the same heading in the
        document or an error is raised.
        """
        id = attr['id']
        title = attr['title']
        if id is not None:
            element = state.references['par', 'id', id]
            if title is not None:
                element2 = state.references['par', 'title', title]
                if element is not element2:
                    state.log(logging.ERROR, 'Mismatch between '
                              'title %r and id %r', title, id)
                    raise ValueError(f'Mismatch between title {title!r} and '
                                     'id {id!r}')
                state.log(logging.INFO, 'Resolved <%s id="%s" title="%s"> '
                          'to "%s"', name, id, title, element)
            else:
                state.log(logging.INFO, 'Resolved <%s id="%s"> to "%s"',
                          name, id, element)
        elif title is not None:
            element = state.references['par', 'title', title]
            state.log(logging.INFO, 'Resolved <%s title="%s"> to "%s"',
                      name, title, element)
        else:
            state.log(logging.ERROR, '<%s> must be identified '
                      'by either title or ID', name)
            raise ValueError('Either "id" or "title" attribute must '
                             f'be present in <{name}>')

        return element


def _register_builtins():
    """
    Called on module import to register all of the
    :ref:`tag-api-builtins`.

    The current implementation registers all subclasses of
    :py:class:`BuiltinTag` whose names end with ``'Tag'``. Tag names are
    extracted from class names by converting CamelCase to camel-case and
    removing the trailing ``'-tag'``. For example,
    :py:class:`TableTag` gets registered for :ref:`xml-spec-tags-table`.
    """
    def name2tag(name):
        """
        Convert CamelCaseNames to all-lowercase-hyphenated-names.
        """
        return name[0].lower() + ''.join(
            c if c.islower() else '-' + c.lower() for c in name[1:]
        )

    from sys import modules
    self = modules[__name__]
    for name in dir(self):
        if name.endswith('Tag'):
            attr = getattr(self, name)
            if isinstance(attr, type) and attr is not BuiltinTag and \
                    issubclass(attr, BuiltinTag):
                name = name2tag(name)[:-4]
                tag_registry[name] = attr()


_register_builtins()
