# -*- coding: utf-8 -*-

# imprint: a program for creating documents from data and content templates
#
# Copyright (C) 2019  Joseph R. Fox-Rabinovitz <jfoxrabinovitz at gmail dot com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# Author: Joseph Fox-Rabinovitz <jfoxrabinovitz at gmail dot com>
# Version: 13 Apr 2019: Initial Coding


"""
Root package for :ref:`plugins-handlers` for inserting tables into a
document.

All the handlers in this module are compatible with the
:ref:`plugin interface <plugins-tables-signature>` used by the
:ref:`xml-spec-tags-table` tag. This package exposes all the handlers
defined in its submodules.
"""

from itertools import repeat

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Emu

from haggis.files.docx import merge_row, set_row_height, str2length


__all__ = ['fixed_width_table',]


#: Mapping of valid alignment values for :py:func:fixed_width_table` to
#: The corresponding :py:class:`~docx.enum.table.WD_TABLE_ALIGNMENT`
#: values. The keys of this mapping are valid values for the
#: ``alignment`` key used in the :ref:`plugins-data-configuration` of
#: the :py:func:`~imprint.core.handler.table.tables.CSVFile` and
#: :py:func:`~imprint.core.handler.table.tables.DataFrame` handlers.
_table_alignment = {
    'left': WD_TABLE_ALIGNMENT.LEFT,
    'right': WD_TABLE_ALIGNMENT.RIGHT,
    'center': WD_TABLE_ALIGNMENT.CENTER,
    '<' : WD_TABLE_ALIGNMENT.LEFT,
    '>': WD_TABLE_ALIGNMENT.RIGHT,
    '^': WD_TABLE_ALIGNMENT.CENTER,
}


def fixed_width_table(doc, style, rows=1, cols=None, alignment='^',
                      width=None, col_widths=None, row_height=None):
    """
    Construct a table of fixed width, with optionally specified column
    widths.

    Parameters
    ----------
    doc : docx.document.Document
        The document to insert the table into.
    style : str
        The name of the style to assign to the document.
    rows : int
        The number of rows to add to the table.
    cols : int or None
        The number of colums to add to the table. If :py:obj:`None`, the
        length of ``col_widths`` is used. Either this argument or
        ``col_widths`` must be non-falsy.
    alignment : { '^', '<', '>' }
        The direction in which the table shoule be aligned. Defaults to
        center (``'^'``).
    width : length-spec
        Either a number in inches, or a string representing a number
        with optional units. A full list of accepted units is provided
        in :py:func:`haggis.files.docx.str2length`. Ignored if
        ``col_widths`` is supplied in a way that does not leave
        unspecified columns. Defaults to 7 inches.
    col_widths : sequence[length-spec]
        A sequence of widths, each in the same format as ``width``, for
        the individual columns of the table. If this is longer than
        ``cols``, only the required number of elements is used. If
        shorter, the remaining columns are divided evenly to fill
        ``width``.
    row_height : None or length-spec or sequence[length-spec]
        The height to assign to each row. :py:obj:`None` is the default,
        meaning to let the rows adjust themselves. If a single number or
        numerical string with optional units, it will be applied to all
        the newly created rows. If a sequence, only the minimum of
        ``rows`` and the length of the sequence will be affected. A
        sequence may contain any of the values allowed as scalars
        (including :py:obj:`None`).

    Return
    ------
    table : docx.table.Table
        The newly created table object.
    """
    if width is None:
        width = 7.0
    width = str2length(width)

    if col_widths is None:
        col_widths = []
    else:
        col_widths = [str2length(w) for w in col_widths]

    if not cols:
        if not col_widths:
            raise('Either the number of columns or their widths must be set')
        cols = len(col_widths)

    table = doc.add_table(rows=rows, cols=0, style=style)
    try:
        table.alignment = _table_alignment[alignment]
    except KeyError:
        raise ValueError(f'Illegal value of alignment {alignment!r}')
    table.autofit = not col_widths

    if cols > len(col_widths):
        extra_cols = cols - len(col_widths)
        total = Emu(sum(w.emu for w in col_widths))
        if total > width:
            raise ValueError(f'Existing columns occupy {total.inches}" '
                             f'of {total.width}".')
        delta = width.emu - total.emu
        col_widths += [Emu(delta / extra_cols)] * (extra_cols - 1) + \
                      [Emu(delta / extra_cols + delta % extra_cols)]
    else:
        col_widths = col_widths[:cols]

    for w in col_widths:
        table.add_column(w)

    try:
        iter(row_height)
    except:
        row_height = repeat(row_height)
    else:
        if isinstance(row_height, str):
            row_height = repeat(row_height)

    for row, height in zip(table.rows, row_height):
        if height is None:
            continue
        height = str2length(height)
        set_row_height(row, height)

    return table


from .tables import *
from .tables import __all__ as _all
__all__.extend(_all)
del _all
